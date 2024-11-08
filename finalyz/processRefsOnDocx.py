"""
a python script to process the references of a docx document

usage:
    python ref_script.py your_document.docx

dependencies:
    relies on the `python-docx` module,
    see: https://github.com/python-openxml/python-docx
"""
import sys
import numpy as np
from docx import Document

filename = sys.argv[1]

document = Document(filename)

########################################################
### find all references in the "References" section ####
########################################################

References = {'short':[],
              'number':[],
              'location':[],
              'description':[]}

# find the reference section
iRefStart = np.flatnonzero([p.text[:10]=='References' for p in document.paragraphs])

if len(iRefStart)>0:
    iRefStart = iRefStart[0]+1
    i = iRefStart
    while (len(document.paragraphs[i].text)>0) and\
            (document.paragraphs[i].text[0]=='['):
        i+=1
    iRefEnd = i
    # now extract all references
    for paragraph in document.paragraphs[iRefStart:iRefEnd]:
        References['short'].append(paragraph.text.split(']')[0].replace('[',''))
        References['description'].append(paragraph.text.split('] ')[1])

else:
    print('')
    print(' [!!] "References" section not found [!!] ')
    print('')


########################################################
### find the reference location and number them ########
########################################################

References['location'] = np.ones(len(References['short']))*np.inf
PSHIFT = 100000

for r, ref in enumerate(References['short']):
    for p, paragraph in enumerate(document.paragraphs):
        if p not in range(iRefStart, iRefEnd+1):
            splits = paragraph.text.split(ref)
            if (References['location'][r]==np.inf) and (len(splits)>1):
                print(ref)
                References['location'][r] = PSHIFT*p+len(splits[0])

# numbered according to increasing position
References['number'] = np.argsort(References['location'])+1

########################################################
### replace the Reference within the text       ########
########################################################

for p, paragraph in enumerate(document.paragraphs):

    # if not in reference section
    if p not in range(iRefStart, iRefEnd+1):

        for r, ref in enumerate(References['short']):

            if np.isfinite(References['location'][r]) and (ref in paragraph.text):

                for run in paragraph.runs:

                    # deal with all situations

                    if ' (%s)' % ref in run.text:
                        # 1) ref alone
                        run.text = run.text.replace(' (%s)' % ref,
                                                    ' [%i]'% References['number'][r])
                        # run.font.superscript = True

                    elif ' (%s; ' % ref in run.text:
                        print(run.text)
                        # 2) ref start with other
                        run.text = run.text.replace(' (%s; ' % ref,
                                                    ' [%i,'% References['number'][r])
                        # run.font.superscript = True

                    elif '%s; ' % ref in run.text:
                        # 3) middle ones
                        run.text = run.text.replace('%s; ' % ref,
                                                    '%i,'% References['number'][r])

                    elif '%s)' % ref in run.text:
                        # 4) ref end with other
                        run.text = run.text.replace('%s)' % ref,
                                                    '%i]'% References['number'][r])

                    # elif ref in run.text:
                    # 5) in text reference
                    run.text = run.text.replace('%s' % ref,
                                                '[%i]'% References['number'][r])

########################################################
### replace references in the "References" section  ####
########################################################

# for p, paragraph in enumerate(document.paragraphs[iRefStart:iRefEnd+1]):
    # i0 = np.flatnonzero(References['number']==(p+1))[0]
    # if np.isfinite(References['location'][i0]):
        # paragraph.text = '[%i] %s' % (References['number'][i0],
                                      # References['description'][i0])
    # else:
        # paragraph.text = '[UNUSED] %s' % References['description'][i0]

document.save(filename.replace('.docx', '-new.docx'))
